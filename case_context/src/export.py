"""Document export functionality module."""

import logging
from pathlib import Path
from typing import Optional
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

logger = logging.getLogger(__name__)


def export_to_docx(
    answer: str, output_path: Optional[Path] = None, title: str = "Strategic Analysis"
) -> Path:
    """
    Export the generated answer to a DOCX file.

    Args:
        answer: The generated answer text
        output_path: Optional path for the output file
        title: Title for the document

    Returns:
        Path to the created document
    """
    if output_path is None:
        output_path = Path(f"{title.lower().replace(' ', '_')}.docx")

    doc = Document()

    # Add title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add answer
    answer_para = doc.add_paragraph()
    answer_para.add_run(answer)

    try:
        doc.save(output_path)
        logger.info(f"Exported document to {output_path}")
        return output_path
    except Exception as e:
        logger.error(f"Failed to export document: {e}")
        raise
