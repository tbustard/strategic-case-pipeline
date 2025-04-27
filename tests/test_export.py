"""Tests for the export module."""

import pytest
from pathlib import Path
from case_context.export import export_to_markdown, export_to_docx
from docx import Document


def test_export_to_markdown(tmp_path):
    """Test exporting to markdown file."""
    out = tmp_path / "out.md"
    export_to_markdown("Hello, world!", str(out))
    assert out.read_text(encoding="utf-8") == "Hello, world!"


def test_export_to_docx(tmp_path):
    """Test exporting to docx file."""
    out = tmp_path / "out.docx"
    export_to_docx("Hello, DOCX!", str(out), overwrite=True, add_heading=False)
    doc = Document(str(out))
    texts = [p.text for p in doc.paragraphs]
    assert "Hello, DOCX!" in texts


def test_export_to_markdown_with_unicode(tmp_path):
    """Test exporting unicode characters to markdown."""
    out = tmp_path / "unicode.md"
    text = "Hello, 世界!"
    export_to_markdown(text, str(out))
    assert out.read_text(encoding="utf-8") == text


def test_export_to_docx_with_unicode(tmp_path):
    """Test exporting unicode characters to docx."""
    out = tmp_path / "unicode.docx"
    text = "Hello, 世界!"
    export_to_docx(text, str(out), add_heading=False)
    doc = Document(str(out))
    texts = [p.text for p in doc.paragraphs]
    assert text in texts


def test_export_to_markdown_empty(tmp_path):
    """Test exporting empty string to markdown."""
    out = tmp_path / "empty.md"
    export_to_markdown("", str(out))
    assert out.read_text(encoding="utf-8") == ""


def test_export_to_docx_empty(tmp_path):
    """Test exporting empty string to docx."""
    out = tmp_path / "empty.docx"
    export_to_docx("", str(out), add_heading=False)
    doc = Document(str(out))
    assert len(doc.paragraphs) == 0  # heading disabled, empty content
 