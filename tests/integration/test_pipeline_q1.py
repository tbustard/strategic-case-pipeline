"""Integration tests for the pipeline."""

import re
import pytest
from pathlib import Path
from docx import Document
from case_context.pipeline import run_pipeline


def _docx_to_text(p: str) -> str:
    """Convert docx file to text."""
    doc = Document(p)
    return "\n".join(par.text for par in doc.paragraphs)


def _norm(t: str) -> str:
    """Normalize text for comparison."""
    return re.sub(r"\s+", " ", t).strip().lower()


@pytest.fixture
def test_files(tmp_path):
    """Create test files for integration testing."""
    # Create test case file
    case_path = tmp_path / "test_case.docx"
    case_doc = Document()
    case_doc.add_paragraph(
        "This is a test case about Network Effects and First Mover Advantage."
    )
    case_doc.save(case_path)

    # Create test question file
    question_path = tmp_path / "test_question.docx"
    question_doc = Document()
    question_doc.add_paragraph("What strategic concepts are relevant to this case?")
    question_doc.save(question_path)

    # Create test instructions file
    instructions_path = tmp_path / "test_instructions.docx"
    instructions_doc = Document()
    instructions_doc.add_paragraph("Analyze the case and identify strategic concepts.")
    instructions_doc.save(instructions_path)

    # Create test style file
    style_path = tmp_path / "test_style.docx"
    style_doc = Document()
    style_doc.add_paragraph("Write in a clear, academic style.")
    style_doc.save(style_path)

    return {
        "case_path": str(case_path),
        "question_path": str(question_path),
        "instructions_path": str(instructions_path),
        "style_path": str(style_path),
    }


@pytest.mark.slow
def test_question1_full_flow(test_files):
    """Test the complete pipeline for question 1."""
    out = run_pipeline(**test_files)

    # Check section headers
    assert "strategic concepts" in out.lower()
    assert "case-specific evidence" in out.lower()
    assert "additional context" in out.lower()

    # Check for expected content
    assert "network effects" in out.lower()
    assert "first mover advantage" in out.lower()
