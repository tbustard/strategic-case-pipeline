"""Extended tests for the export module."""

import pytest
from pathlib import Path
from case_context.export import export_to_markdown, export_to_docx


def test_markdown_overwrite(tmp_path):
    """Test markdown export with overwrite handling."""
    out = tmp_path / "dup.md"
    out.write_text("old", encoding="utf-8")

    # Test that overwrite=False raises error
    with pytest.raises(FileExistsError):
        export_to_markdown("new", str(out))

    # Test that overwrite=True works
    export_to_markdown("new", str(out), overwrite=True)
    assert out.read_text(encoding="utf-8") == "new"


def test_docx_overwrite(tmp_path):
    """Test docx export with overwrite handling."""
    out = tmp_path / "dup.docx"

    # Create initial file
    export_to_docx("old", str(out), overwrite=True)

    # Test that overwrite=False raises error
    with pytest.raises(FileExistsError):
        export_to_docx("new", str(out))

    # Test that overwrite=True works
    export_to_docx("new", str(out), overwrite=True)

    # Verify content
    from docx import Document

    doc = Document(str(out))
    assert doc.paragraphs[0].text == "new"


def test_markdown_line_handling(tmp_path):
    """Test markdown export handles multi-line content."""
    out = tmp_path / "lines.md"
    content = "Line 1\nLine 2\nLine 3"
    export_to_markdown(content, str(out))
    assert out.read_text(encoding="utf-8") == content


def test_docx_line_handling(tmp_path):
    """Test docx export handles multi-line content."""
    out = tmp_path / "lines.docx"
    content = "Line 1\nLine 2\nLine 3"
    export_to_docx(content, str(out))

    from docx import Document

    doc = Document(str(out))
    assert len(doc.paragraphs) == 3
    assert doc.paragraphs[0].text == "Line 1"
    assert doc.paragraphs[1].text == "Line 2"
    assert doc.paragraphs[2].text == "Line 3"
