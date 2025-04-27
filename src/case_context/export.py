"""Export functionality for case analysis answers."""

from pathlib import Path
from typing import Optional
from docx import Document
from docx.shared import Inches, Pt
import os


def export_to_markdown(answer: str, path: str, overwrite: bool = False) -> None:
    """Write the answer string to a Markdown file at `path`.

    Args:
        answer: The answer text to export
        path: Path where the markdown file should be saved
        overwrite: If True, overwrite existing file. If False, raise FileExistsError.

    Raises:
        FileExistsError: If file exists and overwrite=False
    """
    p = Path(path)
    if p.exists() and not overwrite:
        raise FileExistsError(f"{path} already exists. Use overwrite=True to replace.")
    p.write_text(answer, encoding="utf-8")


def export_to_docx(text: str, output_path: str, overwrite: bool = False, add_heading: bool = True) -> None:
    """Export text to a .docx file.

    Args:
        text: The text content to export
        output_path: Path to save the .docx file
        overwrite: Whether to overwrite existing file
        add_heading: Whether to add a heading to the document
    """
    if not overwrite and os.path.exists(output_path):
        raise FileExistsError(f"File {output_path} already exists")

    doc = Document()
    
    # Split text into lines and create paragraphs
    lines = text.split("\n")
    for i, line in enumerate(lines):
        if line.strip():  # Only add non-empty lines
            if i == 0 and not add_heading:
                # First line should be normal text if add_heading is False
                doc.add_paragraph(line)
            elif line.startswith("#"):
                # Handle markdown-style headings
                level = line.count("#")
                text = line.lstrip("#").strip()
                doc.add_heading(text, level)
            else:
                doc.add_paragraph(line)

    doc.save(output_path)


def export_to_txt(text: str, output_path: str) -> None:
    """Export text to a .txt file.

    Args:
        text: Text content to export
        output_path: Path to save the .txt file
    """
    Path(output_path).write_text(text)
